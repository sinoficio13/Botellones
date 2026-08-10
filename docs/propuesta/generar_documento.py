# -*- coding: utf-8 -*-
"""Genera la propuesta comercial (Word) para el cliente."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AZUL     = RGBColor(0x13, 0x5E, 0xA8)
AZUL_OSC = RGBColor(0x0B, 0x3A, 0x6E)
AZUL_BG  = 'E7F0FA'
VERDE    = RGBColor(0x1E, 0x8E, 0x3E)
GRIS     = RGBColor(0x4B, 0x55, 0x63)
TEXTO    = RGBColor(0x26, 0x2B, 0x33)
NARANJA  = 'FFF4E5'

doc = Document()

sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.2)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = TEXTO
normal.paragraph_format.space_after = Pt(6)


def shade_paragraph(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def add_bar(text, fill='135EA8', color=RGBColor(0xFF, 0xFF, 0xFF), size=13):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(10)
    shade_paragraph(p, fill)
    r = p.add_run('  ' + text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def add_sub(text, color=AZUL, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def add_text(text, size=11, color=None, italic=False, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if italic:
        r.italic = True
    return p


def add_bullet(text, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_lead:
        r = p.add_run(bold_lead + ': ')
        r.bold = True
        r.font.color.rgb = AZUL_OSC
    p.add_run(text)
    return p


def add_check(text):
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run('✔  ')
    r1.bold = True
    r1.font.color.rgb = VERDE
    p.add_run(text)
    return p


def add_step(n, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    r = p.add_run(str(n) + '.  ')
    r.bold = True
    r.font.color.rgb = AZUL
    p.add_run(text)
    return p


def add_box(title, lines, fill=AZUL_BG, border='135EA8'):
    t = doc.add_table(rows=1, cols=1)
    tbl = t._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:' + edge)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '10')
        el.set(qn('w:color'), border)
        borders.append(el)
    tblPr.append(borders)
    cell = t.cell(0, 0)
    shade_cell(cell, fill)
    if title:
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = AZUL_OSC
    for line in lines:
        p = cell.add_paragraph()
        r = p.add_run(line)
        r.font.color.rgb = AZUL_OSC
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def add_hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '135EA8')
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'PAGE'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    run._r.append(f1); run._r.append(it); run._r.append(f2)
    run.font.size = Pt(9)
    run.font.color.rgb = GRIS


# ------------------------------------------------------------------ PORTADA
for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Distribuidora de Agua')
r.font.size = Pt(16); r.bold = True; r.font.color.rgb = AZUL_OSC

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Sistema de Gestión de Botellones')
r.font.size = Pt(30); r.bold = True; r.font.color.rgb = AZUL

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Propuesta de sistema digital para el control de clientes,\nbotellones y recargas')
r.font.size = Pt(14); r.font.color.rgb = GRIS

doc.add_paragraph()
add_hr()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Documento para revisión y debate  •  Agosto 2026')
r.font.size = Pt(12); r.font.color.rgb = GRIS

doc.add_page_break()

# ------------------------------------------------------------------ INTRO
add_bar('¿QUÉ ES ESTA PROPUESTA?')
add_text('Un sistema digital pensado para el día a día de una distribuidora de agua. '
         'Se usa desde el celular del repartidor y desde la computadora de la oficina, '
         'y guarda de forma ordenada toda la información del negocio: clientes, '
         'botellones y recargas.')
add_text('La idea es simple: cada botellón, cada cliente y cada entrega quedan '
         'registrados de una vez, y la información está siempre a mano, '
         'sin papeles y sin planillas sueltas.', italic=True)

# ------------------------------------------------------------------ PROBLEMA
add_bar('¿QUÉ PROBLEMA RESUELVE?')
add_text('Con este sistema vas a poder:')
add_check('Tener todos los clientes registrados con sus datos, dirección, fotos y preferencias')
add_check('Saber en todo momento cuántos botellones tiene cada cliente y en qué estado están')
add_check('Registrar cada recarga con fecha, hora y quién la realizó')
add_check('Que el repartidor llegue sin perderse: mapa + foto de la fachada + referencias')
add_check('Ver el historial completo de cada cliente y de cada botellón')
add_check('Saber qué clientes están pendientes de visita y quiénes consumen más')

# ------------------------------------------------------------------ DIA A DIA
add_bar('ASÍ TRABAJA EL SISTEMA EN EL DÍA A DÍA')
add_box('Un día con el sistema', [
    '1.  Por la mañana, el repartidor abre el mapa en el celular y arma la ruta del día con todos los clientes ubicados.',
    '2.  En cada parada, la ficha del cliente muestra la foto de la fachada, la entrada y la referencia ("casa azul detrás de la farmacia San José").',
    '3.  Al entregar, toca "Registrar recarga", elige el botellón y confirma. Se toma menos de un minuto.',
    '4.  En la oficina, el administrador ve el resumen del día: cuántas recargas se hicieron, qué clientes quedaron pendientes y qué botellones faltan en la planta.',
])

# ------------------------------------------------------------------ MODULOS
add_bar('LOS MÓDULOS DEL SISTEMA')

add_sub('1.  Registro de clientes')
add_bullet('ficha completa: nombre, negocio, teléfonos y WhatsApp', bold_lead='Ficha del cliente')
add_bullet('tipo de cliente: Casa, Negocio, Oficina u Otro', bold_lead='Tipo')
add_bullet('código único (CL-0001) para identificarlo rápido', bold_lead='Código')

add_sub('2.  Dirección y ubicación (GPS)')
add_bullet('calle, sector, ciudad, referencia de ubicación', bold_lead='Dirección escrita')
add_bullet('la ubicación GPS que llega por WhatsApp se guarda automáticamente en la ficha', bold_lead='GPS')
add_bullet('botón "Ver en Google Maps" que abre la ruta desde donde esté el repartidor', bold_lead='Navegación')

add_sub('3.  Fotografías del cliente')
add_bullet('de la fachada, de la entrada, de referencias cercanas y fotos adicionales', bold_lead='4 tipos de foto')
add_bullet('se toman con la cámara del celular en el momento', bold_lead='Cámara')
add_bullet('ayudan a reconocer el lugar a la distancia y a no equivocarse de casa', bold_lead='Beneficio')

add_sub('4.  Preferencias del cliente')
add_bullet('Mañana, Tarde o Noche', bold_lead='Horario preferido')
add_bullet('días de entrega y forma de contacto preferida', bold_lead='Días')
add_bullet('"golpear fuerte porque no escucha el timbre", "entrar por el portón lateral", "llamar antes de llegar"', bold_lead='Observaciones')

add_sub('5.  Control de botellones')
add_bullet('cada botellón tiene su propio código (BOT-00001)', bold_lead='Código')
add_bullet('Disponible, Asignado a un cliente, En recarga, En mantenimiento, Dañado o Perdido', bold_lead='Estado en todo momento')
add_bullet('se sabe exactamente qué botellones tiene cada cliente', bold_lead='Asignación')

add_sub('6.  Historial de recargas')
add_bullet('cada recarga queda registrada con fecha, hora, botellón y quién la hizo', bold_lead='Registro automático')
add_bullet('historial completo por cliente y por botellón', bold_lead='Historial')
add_bullet('la última recarga de cada cliente se ve de un vistazo', bold_lead='Consulta rápida')

add_sub('7.  Panel de control')
add_bullet('clientes, botellones, recargas de hoy y del mes', bold_lead='Resumen del día')
add_bullet('nuevo cliente, registrar botellón, registrar recarga, buscar y ver el mapa', bold_lead='Accesos rápidos')

add_sub('8.  Búsqueda de clientes')
add_bullet('buscar por nombre, teléfono, código o dirección', bold_lead='Varias formas')
add_bullet('resultados al instante, incluso con datos incompletos', bold_lead='Rapidez')

add_sub('9.  Mapa de clientes')
add_bullet('todos los clientes con su ubicación en un mapa', bold_lead='Mapa')
add_bullet('al tocar un cliente se abre la ruta en Google Maps', bold_lead='Ruta')

add_sub('10.  Control de operaciones')
add_bullet('clientes pendientes de visita', bold_lead='Pendientes')
add_bullet('botellones en recarga o en mantenimiento', bold_lead='Inventario')
add_bullet('clientes con mayor consumo para conocer a los más importantes', bold_lead='Consumo')

# ------------------------------------------------------------------ BENEFICIOS
add_bar('BENEFICIOS CLAVE')
add_check('Ahorro de tiempo en el reparto y en la oficina')
add_check('Menos errores: cada botellón tiene su historial')
add_check('Control total del inventario de botellones')
add_check('Mejor atención: se sabe qué prefiere cada cliente')
add_check('El repartidor nunca se pierde: GPS + fotos + referencias')
add_check('Toda la información respaldada y ordenada, sin papeles')

# ------------------------------------------------------------------ CELULAR
add_bar('CÓMO SE USA')
add_sub('En el celular del repartidor')
add_bullet('se instala como una app o se abre desde el navegador', bold_lead='Sencillo')
add_bullet('funciona con la señal de datos del celular', bold_lead='Datos')
add_bullet('registrar una recarga toma menos de un minuto', bold_lead='Rápido')
add_sub('En la computadora de la oficina')
add_bullet('el administrador gestiona clientes, botellones y revisa el resumen', bold_lead='Gestión')

# ------------------------------------------------------------------ SEGURIDAD
add_bar('SEGURIDAD Y USUARIOS')
add_bullet('cada persona entra con su usuario y contraseña', bold_lead='Acceso')
add_bullet('Administrador (control total) y Repartidor (registra recargas y consulta clientes)', bold_lead='Roles')
add_bullet('cada recarga queda registrada con quién la hizo', bold_lead='Responsabilidad')

# ------------------------------------------------------------------ FUTURO
add_bar('PENSADO PARA CRECER')
add_text('El sistema se diseña desde el inicio para sumar en una próxima etapa:')
add_check('Optimización de rutas de entrega')
add_check('Cobros y pagos por cliente')
add_check('Reportes de consumo y estadísticas')

# ------------------------------------------------------------------ DEBATE
add_bar('PUNTOS PARA DEBATIR JUNTOS')
add_box('Preguntas para la reunión', [
    '¿Cuántas personas van a usar el sistema? (repartidor, administrador, otros)',
    '¿Lo van a usar principalmente desde el celular?',
    '¿Quieren los cobros y pagos desde el inicio o en una próxima etapa?',
    '¿Cómo llega hoy la ubicación de los clientes por WhatsApp?',
], fill=NARANJA, border='C25E0A')

# ------------------------------------------------------------------ CIERRE
add_bar('PRÓXIMOS PASOS')
add_text('Este documento es el punto de partida. La idea es ajustar el sistema a cómo '
         'trabaja la distribuidora hoy, para que se adapte a ustedes y no al revés.')
add_text('Al confirmar los puntos de debate, comenzamos a construir el sistema y '
         'lo ponemos a prueba con datos reales de la operación.')

# ------------------------------------------------------------------ FOOTER
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run('Propuesta Sistema de Gestión de Botellones  •  Página ')
r.font.size = Pt(9); r.font.color.rgb = GRIS
add_page_number_field(fp)

doc.core_properties.title = 'Propuesta Sistema de Gestión de Botellones'
doc.core_properties.author = 'Distribuidora de Agua'

OUT = r'D:\Github\Botellon\docs\Propuesta_Sistema_Botellones.docx'
doc.save(OUT)
print('OK', OUT)
